import docx

file = open('guests.txt')
names = file.readlines()

doc = docx.Document()
for name in names:
    name = name.strip()

    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.add_paragraph(name)
    doc.add_paragraph('at 11010 Memory Lane on the Evening of')                               
    doc.add_paragraph('April 1st')
    doc.add_paragraph("at 7 o'clock")

    doc.add_page_break()

doc.save('custominvite.docx')